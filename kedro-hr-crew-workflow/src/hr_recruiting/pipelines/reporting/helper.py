"""Helper functions for reporting pipeline.

This module contains helper functions used by the reporting pipeline nodes
for email drafting, document generation, and formatting.
"""

import re
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from hr_recruiting.base.utils import get_model_dump
from hr_recruiting.pipelines.reporting.models import EmailDraft
from hr_recruiting.pipelines.screening.models import ScreeningResult


def draft_email(
    screening_result: dict[str, Any],
    email_templates: dict[str, Any],
) -> dict[str, Any]:
    """Draft email communication based on screening result and templates.

    This is a deterministic function that creates an email draft from templates.
    The template is selected based on the recommendation (proceed/review/reject).

    Args:
        screening_result: Screening result with candidate_name, job_title, and recommendation
        email_templates: Email templates dictionary from config, keyed by recommendation type

    Returns:
        EmailDraft dictionary with subject, body, and placeholders
    """
    # Validate screening result first - fail if required fields are missing
    validated_result = ScreeningResult(**screening_result)
    
    # Extract candidate_name and job_title from validated result
    candidate_name = validated_result.candidate_name
    job_title = validated_result.job_title

    # Get recommendation from validated result
    recommendation = validated_result.recommendation.lower()

    # Get template for the recommendation type, default to review if not found
    template = email_templates.get(recommendation, email_templates.get("review", {}))

    # Replace placeholders in template
    subject_template = template.get("subject", "Application Update: {job_title}")
    body_template = template.get("body", "")

    subject = subject_template.format(
        candidate_name=candidate_name,
        job_title=job_title,
    )

    body = body_template.format(
        candidate_name=candidate_name,
        job_title=job_title,
    )

    # Validate using EmailDraft model
    return get_model_dump(EmailDraft,
        subject=subject,
        body=body,
        placeholders={
            "candidate_name": candidate_name,
            "job_title": job_title,
        },
    )


def setup_document(result: ScreeningResult) -> Document:
    """Create and configure the Word document with metadata.
    
    Args:
        result: ScreeningResult to extract application_id for document title
        
    Returns:
        Configured Word document with title, subject, and author metadata
    """
    doc = Document()
    doc.core_properties.title = f"Screening Report - {result.application_id}"
    doc.core_properties.subject = "HR Candidate Screening Report"
    doc.core_properties.author = "HR Recruiting System"
    return doc


def add_header(doc: Document, result: ScreeningResult) -> None:
    """Add document header with title, application ID, and job title.
    
    Args:
        doc: Word document to add header to
        result: ScreeningResult containing candidate_name, application_id, and job_title
    """
    title_text = f"Screening Report for {result.candidate_name}, {result.job_title}"
    title = doc.add_heading(title_text, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Application ID: {result.application_id}")
    doc.add_paragraph("")


def add_summary_table(doc: Document, result: ScreeningResult) -> None:
    """Add executive summary table with match score, recommendation, and coverage.
    
    Args:
        doc: Word document to add table to
        result: ScreeningResult containing match_score, recommendation, and must_have_coverage
    """
    doc.add_heading("Executive Summary", level=1)
    summary_table = doc.add_table(rows=4, cols=2)
    summary_table.style = "Light Grid Accent 1"
    
    status_emoji = {
        "proceed": "✅ Proceed",
        "review": "⚠️ Review",
        "reject": "❌ Reject",
    }
    
    summary_data = [
        ("Match Score", f"{result.match_score:.1f}/100"),
        ("Recommendation", result.recommendation.upper()),
        ("Must-Have Coverage", f"{result.must_have_coverage:.0%}"),
        ("Status", status_emoji.get(result.recommendation, "❓ Unknown")),
    ]
    
    for i, (label, value) in enumerate(summary_data):
        label_cell = summary_table.rows[i].cells[0]
        value_cell = summary_table.rows[i].cells[1]
        label_cell.text = label
        value_cell.text = value
        label_cell.paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph("")


def add_bullet_list_section(
    doc: Document,
    heading: str,
    items: list[str],
    prefix: str = "•",
) -> None:
    """Add a section with heading and bullet list.
    
    Args:
        doc: Word document to add section to
        heading: Section heading text
        items: List of items to display as bullets
        prefix: Prefix character for items (default: "•", can be emoji like "⚠️")
        
    Note:
        If items list is empty, the section is not added.
        Existing bullet characters in items are stripped before adding.
    """
    if not items:
        return
    
    doc.add_heading(heading, level=1)
    for item in items:
        # Strip any existing bullets/whitespace from the item
        # Remove common bullet characters and whitespace from the start
        clean_item = item.strip()
        # Remove bullet characters (•, -, *, and various Unicode bullets) from the start
        clean_item = re.sub(r'^[\s•\-\*\u2022\u2023\u2043\u204C\u204D\u2219]+', '', clean_item).strip()
        
        # If prefix is a bullet character, don't add it (List Bullet style adds its own)
        # For other prefixes (like emojis), add them before the item
        if prefix == "•":
            doc.add_paragraph(clean_item, style="List Bullet")
        else:
            doc.add_paragraph(f"{prefix} {clean_item}", style="List Bullet")
    doc.add_paragraph("")


def add_match_results_table(doc: Document, match_results: list[Any]) -> None:
    """Add detailed match results table showing requirement matches.
    
    Args:
        doc: Word document to add table to
        match_results: List of MatchResult objects with requirement, confidence, and snippet_ids
        
    Note:
        If match_results is empty, the table is not added.
    """
    if not match_results:
        return
    
    doc.add_heading("Detailed Requirement Matches", level=1)
    match_table = doc.add_table(rows=len(match_results) + 1, cols=3)
    match_table.style = "Light Grid Accent 1"
    
    # Header row
    headers = ["Requirement", "Confidence", "Evidence Count"]
    header_cells = match_table.rows[0].cells
    for i, header_text in enumerate(headers):
        header_cells[i].text = header_text
        header_cells[i].paragraphs[0].runs[0].bold = True
    
    # Data rows
    for i, match in enumerate(match_results, start=1):
        row_cells = match_table.rows[i].cells
        row_cells[0].text = match.requirement
        row_cells[1].text = f"{match.confidence:.0%}"
        row_cells[2].text = str(len(match.snippet_ids))
    
    doc.add_paragraph("")


def add_email_draft_section(doc: Document, email_draft: Any) -> None:
    """Add email draft section with subject and body.
    
    Args:
        doc: Word document to add section to
        email_draft: EmailDraft object with subject and body
        
    Note:
        If email_draft is None or empty, the section is not added.
    """
    if not email_draft:
        return
    
    doc.add_heading("Drafted Communication", level=1)
    doc.add_paragraph(f"Subject: {email_draft.subject}")
    doc.add_paragraph("")
    doc.add_paragraph("Body:")
    doc.add_paragraph(email_draft.body)
    doc.add_paragraph("")


def add_footer(doc: Document) -> None:
    """Add document footer with system information.
    
    Args:
        doc: Word document to add footer to
    """
    doc.add_paragraph("─" * 50)
    doc.add_paragraph("This report was generated automatically by the HR Recruiting System.")
    doc.add_paragraph("Please review all recommendations before taking action.")
