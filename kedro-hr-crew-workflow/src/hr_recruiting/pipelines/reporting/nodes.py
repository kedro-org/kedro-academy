"""Reporting pipeline nodes - deterministic functions."""

from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from hr_recruiting.pipelines.screening.models import ScreeningResult


def validate_screening_result(screening_result_raw: dict[str, Any]) -> dict[str, Any]:
    """Validate screening result.

    Args:
        screening_result_raw: Raw screening result dictionary

    Returns:
        Validated screening result dictionary
    """
    try:
        validated = ScreeningResult(**screening_result_raw)
        return validated.model_dump()
    except Exception as e:
        raise ValueError(f"Screening result validation failed: {e}") from e


def _setup_document(result: ScreeningResult) -> Document:
    """Create and configure the Word document.

    Args:
        result: Screening result model

    Returns:
        Configured Document object
    """
    doc = Document()
    doc.core_properties.title = f"Screening Report - {result.application_id}"
    doc.core_properties.subject = "HR Candidate Screening Report"
    doc.core_properties.author = "HR Recruiting System"
    return doc


def _add_header(doc: Document, application_id: str) -> None:
    """Add document header with title and application ID.

    Args:
        doc: Document to add header to
        application_id: Application identifier
    """
    title = doc.add_heading("Candidate Screening Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Application ID: {application_id}")
    doc.add_paragraph("")


def _add_summary_table(doc: Document, result: ScreeningResult) -> None:
    """Add executive summary table.

    Args:
        doc: Document to add table to
        result: Screening result model
    """
    doc.add_heading("Executive Summary", level=1)
    summary_table = doc.add_table(rows=4, cols=2)
    summary_table.style = "Light Grid Accent 1"
    
    status_emoji = {
        "proceed": "✅ Proceed",
        "review": "⚠️ Review",
        "reject": "❌ Reject",
    }
    
    summary_data = [
        ("Match Score", f"{result.match_score:.1f}/100"),
        ("Recommendation", result.recommendation.upper()),
        ("Must-Have Coverage", f"{result.must_have_coverage:.0%}"),
        ("Status", status_emoji.get(result.recommendation, "❓ Unknown")),
    ]
    
    for i, (label, value) in enumerate(summary_data):
        label_cell = summary_table.rows[i].cells[0]
        value_cell = summary_table.rows[i].cells[1]
        label_cell.text = label
        value_cell.text = value
        label_cell.paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph("")


def _add_bullet_list_section(
    doc: Document,
    heading: str,
    items: list[str],
    prefix: str = "•",
) -> None:
    """Add a section with heading and bullet list.

    Args:
        doc: Document to add section to
        heading: Section heading text
        items: List of items to display
        prefix: Prefix for each item (default: bullet point)
    """
    if not items:
        return
    
    doc.add_heading(heading, level=1)
    for item in items:
        doc.add_paragraph(f"{prefix} {item}", style="List Bullet")
    doc.add_paragraph("")


def _add_match_results_table(doc: Document, match_results: list[Any]) -> None:
    """Add detailed match results table.

    Args:
        doc: Document to add table to
        match_results: List of match result models
    """
    if not match_results:
        return
    
    doc.add_heading("Detailed Requirement Matches", level=1)
    match_table = doc.add_table(rows=len(match_results) + 1, cols=3)
    match_table.style = "Light Grid Accent 1"
    
    # Header row
    headers = ["Requirement", "Confidence", "Evidence Count"]
    header_cells = match_table.rows[0].cells
    for i, header_text in enumerate(headers):
        header_cells[i].text = header_text
        header_cells[i].paragraphs[0].runs[0].bold = True
    
    # Data rows
    for i, match in enumerate(match_results, start=1):
        row_cells = match_table.rows[i].cells
        row_cells[0].text = match.requirement
        row_cells[1].text = f"{match.confidence:.0%}"
        row_cells[2].text = str(len(match.snippet_ids))
    
    doc.add_paragraph("")


def _add_email_draft_section(doc: Document, email_draft: Any) -> None:
    """Add email draft section.

    Args:
        doc: Document to add section to
        email_draft: Email draft model
    """
    if not email_draft:
        return
    
    doc.add_heading("Drafted Communication", level=1)
    doc.add_paragraph(f"Subject: {email_draft.subject}")
    doc.add_paragraph("")
    doc.add_paragraph("Body:")
    doc.add_paragraph(email_draft.body)
    doc.add_paragraph("")


def _add_footer(doc: Document) -> None:
    """Add document footer.

    Args:
        doc: Document to add footer to
    """
    doc.add_paragraph("─" * 50)
    doc.add_paragraph("This report was generated automatically by the HR Recruiting System.")
    doc.add_paragraph("Please review all recommendations before taking action.")


def generate_hr_report(screening_result: dict[str, Any]) -> Document:
    """Generate HR report from screening result.

    This is a deterministic node that creates a human-readable Word document report.

    Args:
        screening_result: Validated screening result dictionary

    Returns:
        Word document (python-docx Document) with formatted HR report
    """
    result = ScreeningResult(**screening_result)

    # Setup document
    doc = _setup_document(result)
    
    # Add header
    _add_header(doc, result.application_id)
    
    # Add executive summary
    _add_summary_table(doc, result)
    
    # Add candidate strengths
    _add_bullet_list_section(doc, "Candidate Strengths", result.strengths)
    
    # Add identified gaps
    _add_bullet_list_section(doc, "Identified Gaps", result.gaps)
    
    # Add risk flags
    _add_bullet_list_section(doc, "Risk Flags", result.risk_flags, prefix="⚠️")
    
    # Add detailed match results
    _add_match_results_table(doc, result.match_results)
    
    # Add email draft
    _add_email_draft_section(doc, result.email_draft)
    
    # Add QA suggestions
    _add_bullet_list_section(doc, "Next Steps / QA Suggestions", result.qa_suggestions)
    
    # Add footer
    _add_footer(doc)
    
    return doc
