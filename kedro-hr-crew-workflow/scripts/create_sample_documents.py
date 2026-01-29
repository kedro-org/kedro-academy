"""Script to create sample Word documents for testing the HR recruiting pipeline.

This script generates:
- 1 job posting for a Software Engineer (backend focused)
- 1 resume for a Data Scientist
- 1 resume for a Software Engineer
"""

from pathlib import Path

from docx import Document


def create_software_engineer_job_posting(output_path: Path) -> None:
    """Create a job posting for a Software Engineer (backend focused).

    Args:
        output_path: Path where the document should be saved
    """
    doc = Document()
    
    # Set document properties
    doc.core_properties.title = "SE_BACKEND_001"
    doc.core_properties.subject = "Senior Software Engineer - Backend"
    doc.core_properties.author = "HR Recruiting Team"
    
    # Add title
    doc.add_heading("Senior Software Engineer - Backend", 0)
    
    # Company and location
    doc.add_paragraph("Company: TechFlow Solutions")
    doc.add_paragraph("Location: San Francisco, CA (Hybrid - 3 days in office)")
    doc.add_paragraph("Employment Type: Full-time")
    doc.add_paragraph("")  # Empty line
    
    # Job description
    doc.add_heading("Job Description", level=1)
    doc.add_paragraph(
        "We are seeking an experienced Senior Software Engineer to join our backend engineering team. "
        "You will be responsible for designing, developing, and maintaining scalable backend systems "
        "that power our platform serving millions of users. This role requires strong technical skills, "
        "leadership capabilities, and a passion for building robust, high-performance systems."
    )
    doc.add_paragraph("")  # Empty line
    
    # Key responsibilities
    doc.add_heading("Key Responsibilities", level=1)
    responsibilities = [
        "Design and develop scalable backend services using Python, Java, and Go",
        "Build and maintain RESTful APIs and microservices architecture",
        "Collaborate with cross-functional teams including product, design, and frontend engineers",
        "Optimize database queries and improve system performance",
        "Participate in code reviews and maintain high code quality standards",
        "Mentor junior engineers and contribute to technical decision-making",
        "Implement CI/CD pipelines and ensure reliable deployments",
        "Troubleshoot and resolve production issues in a timely manner"
    ]
    for resp in responsibilities:
        doc.add_paragraph(resp, style="List Bullet")
    doc.add_paragraph("")  # Empty line
    
    # Must-have requirements
    doc.add_heading("Must-Have Requirements", level=1)
    must_haves = [
        "5+ years of professional software development experience",
        "Strong proficiency in Python, Java, or Go",
        "Experience with RESTful API design and development",
        "Solid understanding of relational databases (PostgreSQL, MySQL) and NoSQL databases (MongoDB, Redis)",
        "Experience with cloud platforms (AWS, GCP, or Azure)",
        "Knowledge of containerization technologies (Docker, Kubernetes)",
        "Experience with version control systems (Git) and CI/CD pipelines",
        "Strong problem-solving skills and attention to detail",
        "Bachelor's degree in Computer Science, Engineering, or related field"
    ]
    for req in must_haves:
        doc.add_paragraph(req, style="List Bullet")
    doc.add_paragraph("")  # Empty line
    
    # Nice-to-have requirements
    doc.add_heading("Nice-to-Have Requirements", level=1)
    nice_to_haves = [
        "Experience with message queues (RabbitMQ, Kafka, SQS)",
        "Knowledge of GraphQL API design",
        "Experience with distributed systems and microservices architecture",
        "Familiarity with monitoring and observability tools (Prometheus, Grafana, Datadog)",
        "Open source contributions",
        "Experience with test-driven development (TDD)",
        "Knowledge of system design patterns and best practices"
    ]
    for req in nice_to_haves:
        doc.add_paragraph(req, style="List Bullet")
    doc.add_paragraph("")  # Empty line
    
    # Compensation and benefits
    doc.add_heading("Compensation & Benefits", level=1)
    doc.add_paragraph("Salary Range: $150,000 - $200,000 (based on experience)")
    doc.add_paragraph("Benefits:")
    benefits = [
        "Comprehensive health, dental, and vision insurance",
        "401(k) with company matching",
        "Flexible PTO and paid holidays",
        "Professional development budget",
        "Stock options",
        "Remote work flexibility"
    ]
    for benefit in benefits:
        doc.add_paragraph(benefit, style="List Bullet")
    
    # Save document
    doc.save(output_path)
    print(f"Created job posting: {output_path}")


def create_data_scientist_resume(output_path: Path) -> None:
    """Create a resume for a Data Scientist.

    Args:
        output_path: Path where the document should be saved
    """
    doc = Document()
    
    # Set document properties
    doc.core_properties.title = "DS_CANDIDATE_001"
    doc.core_properties.subject = "Sarah Chen - Data Scientist Resume"
    doc.core_properties.author = "Sarah Chen"
    
    # Header
    doc.add_heading("Sarah Chen", 0)
    doc.add_paragraph("Data Scientist | Machine Learning Engineer")
    doc.add_paragraph("Email: sarah.chen@email.com | Phone: (555) 123-4567")
    doc.add_paragraph("LinkedIn: linkedin.com/in/sarahchen | GitHub: github.com/sarahchen")
    doc.add_paragraph("")  # Empty line
    
    # Professional summary
    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph(
        "Experienced Data Scientist with 6+ years of expertise in machine learning, statistical analysis, "
        "and data engineering. Proven track record of building predictive models and deploying ML solutions "
        "at scale. Strong background in Python, SQL, and cloud platforms. Passionate about leveraging data "
        "to drive business decisions and solve complex problems."
    )
    doc.add_paragraph("")  # Empty line
    
    # Skills
    doc.add_heading("Technical Skills", level=1)
    doc.add_paragraph("Programming Languages: Python, R, SQL, Scala")
    doc.add_paragraph("ML/AI: Scikit-learn, TensorFlow, PyTorch, XGBoost, NLP, Computer Vision")
    doc.add_paragraph("Data Tools: Pandas, NumPy, Spark, Airflow, Jupyter")
    doc.add_paragraph("Cloud & Infrastructure: AWS (SageMaker, S3, EC2), GCP, Docker, Kubernetes")
    doc.add_paragraph("Databases: PostgreSQL, MongoDB, Redis")
    doc.add_paragraph("Visualization: Tableau, Matplotlib, Seaborn, Plotly")
    doc.add_paragraph("")  # Empty line
    
    # Work experience
    doc.add_heading("Work Experience", level=1)
    
    # Job 1
    doc.add_heading("Senior Data Scientist", level=2)
    doc.add_paragraph("DataCorp Analytics | San Francisco, CA | Jan 2021 - Present")
    doc.add_paragraph("Key Achievements:")
    achievements1 = [
        "Developed and deployed ML models that improved customer retention by 25%",
        "Built end-to-end ML pipelines using Python, Spark, and Airflow processing 10TB+ daily",
        "Led a team of 3 junior data scientists and mentored them on best practices",
        "Created recommendation systems using collaborative filtering and deep learning",
        "Collaborated with engineering teams to productionize models using Docker and Kubernetes"
    ]
    for achievement in achievements1:
        doc.add_paragraph(achievement, style="List Bullet")
    doc.add_paragraph("")  # Empty line
    
    # Job 2
    doc.add_heading("Data Scientist", level=2)
    doc.add_paragraph("TechStart Inc. | Seattle, WA | Jun 2018 - Dec 2020")
    doc.add_paragraph("Key Achievements:")
    achievements2 = [
        "Built predictive models for fraud detection reducing false positives by 30%",
        "Performed statistical analysis and A/B testing to optimize product features",
        "Developed NLP models for sentiment analysis and text classification",
        "Created automated reporting dashboards using Tableau and Python"
    ]
    for achievement in achievements2:
        doc.add_paragraph(achievement, style="List Bullet")
    doc.add_paragraph("")  # Empty line
    
    # Job 3
    doc.add_heading("Junior Data Analyst", level=2)
    doc.add_paragraph("Analytics Solutions | Boston, MA | Jul 2016 - May 2018")
    doc.add_paragraph("Key Achievements:")
    achievements3 = [
        "Performed data analysis and created visualizations for business stakeholders",
        "Cleaned and preprocessed large datasets using Python and SQL",
        "Assisted in building regression models for sales forecasting"
    ]
    for achievement in achievements3:
        doc.add_paragraph(achievement, style="List Bullet")
    doc.add_paragraph("")  # Empty line
    
    # Education
    doc.add_heading("Education", level=1)
    doc.add_heading("Master of Science in Data Science", level=2)
    doc.add_paragraph("Stanford University | Stanford, CA | 2016")
    doc.add_paragraph("Relevant Coursework: Machine Learning, Statistical Methods, Big Data Analytics")
    doc.add_paragraph("")  # Empty line
    
    doc.add_heading("Bachelor of Science in Mathematics", level=2)
    doc.add_paragraph("UC Berkeley | Berkeley, CA | 2014")
    doc.add_paragraph("Minor: Computer Science")
    
    # Save document
    doc.save(output_path)
    print(f"Created resume: {output_path}")


def create_software_engineer_resume(output_path: Path) -> None:
    """Create a resume for a Software Engineer.

    Args:
        output_path: Path where the document should be saved
    """
    doc = Document()
    
    # Set document properties
    doc.core_properties.title = "SE_CANDIDATE_001"
    doc.core_properties.subject = "Michael Rodriguez - Software Engineer Resume"
    doc.core_properties.author = "Michael Rodriguez"
    
    # Header
    doc.add_heading("Michael Rodriguez", 0)
    doc.add_paragraph("Senior Software Engineer | Backend Specialist")
    doc.add_paragraph("Email: michael.rodriguez@email.com | Phone: (555) 987-6543")
    doc.add_paragraph("LinkedIn: linkedin.com/in/mrodriguez | GitHub: github.com/mrodriguez")
    doc.add_paragraph("")  # Empty line
    
    # Professional summary
    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph(
        "Senior Software Engineer with 6+ years of professional software development experience specializing "
        "in backend systems and distributed architectures. Strong proficiency in Python, Java, and Go with "
        "extensive experience in RESTful API design, microservices, and cloud platforms. Proven track record "
        "of building scalable systems, optimizing database performance, and implementing robust CI/CD pipelines. "
        "Passionate about system design patterns, test-driven development, and mentoring junior engineers."
    )
    doc.add_paragraph("")  # Empty line
    
    # Skills
    doc.add_heading("Technical Skills", level=1)
    doc.add_paragraph("Programming Languages: Python, Java, Go, JavaScript, TypeScript")
    doc.add_paragraph("Backend Frameworks: FastAPI, Django, Flask, Spring Boot, Gin")
    doc.add_paragraph("Databases: PostgreSQL, MySQL, MongoDB, Redis, Cassandra")
    doc.add_paragraph("Cloud & Infrastructure: AWS (EC2, S3, Lambda, RDS), GCP (Cloud Run, Cloud SQL), Azure, Docker, Kubernetes")
    doc.add_paragraph("API Design: RESTful APIs, GraphQL, gRPC, API Gateway design patterns")
    doc.add_paragraph("Version Control & CI/CD: Git, GitHub Actions, Jenkins, GitLab CI, automated testing and deployment pipelines")
    doc.add_paragraph("Message Queues: RabbitMQ, Apache Kafka, AWS SQS")
    doc.add_paragraph("Monitoring & Observability: Prometheus, Grafana, Datadog, ELK Stack, distributed tracing")
    doc.add_paragraph("Development Practices: Test-Driven Development (TDD), system design patterns, code reviews, agile methodologies")
    doc.add_paragraph("")  # Empty line
    
    # Work experience
    doc.add_heading("Work Experience", level=1)
    
    # Job 1
    doc.add_heading("Senior Software Engineer", level=2)
    doc.add_paragraph("CloudScale Technologies | San Francisco, CA | Mar 2020 - Present")
    doc.add_paragraph("Key Achievements:")
    achievements1 = [
        "Architected and developed distributed microservices handling 10M+ requests per day using Python, Java, and Go",
        "Designed and implemented RESTful APIs and GraphQL endpoints following industry best practices, serving mobile and web clients",
        "Led migration from monolithic application to microservices architecture, improving scalability by 300%",
        "Optimized PostgreSQL and MySQL database queries, reducing query response times by 50% and implementing Redis caching strategies",
        "Built and maintained NoSQL databases (MongoDB, Redis) for high-performance data storage and retrieval",
        "Deployed and managed applications on AWS (EC2, S3, RDS) and GCP, utilizing Docker and Kubernetes for containerization and orchestration",
        "Implemented comprehensive CI/CD pipelines using Git, Jenkins, and GitHub Actions, reducing deployment time from 2 hours to 15 minutes",
        "Built real-time data processing pipelines using Apache Kafka, RabbitMQ, and AWS SQS for asynchronous message handling",
        "Set up monitoring and observability using Prometheus, Grafana, and Datadog for system health and performance tracking",
        "Practiced Test-Driven Development (TDD) and applied system design patterns to build maintainable and scalable solutions",
        "Mentored 4 junior engineers, conducted code reviews, and established development best practices"
    ]
    for achievement in achievements1:
        doc.add_paragraph(achievement, style="List Bullet")
    doc.add_paragraph("")  # Empty line
    
    # Job 2
    doc.add_heading("Software Engineer", level=2)
    doc.add_paragraph("StartupHub | Austin, TX | Jun 2017 - Feb 2020")
    doc.add_paragraph("Key Achievements:")
    achievements2 = [
        "Developed scalable backend services using Python, Java, and Django for a SaaS platform serving 1M+ users",
        "Designed and developed RESTful APIs following REST principles, ensuring consistent and intuitive API design",
        "Built and maintained PostgreSQL and MySQL databases, optimized complex queries, and implemented database indexing strategies",
        "Worked with MongoDB and Redis for caching and session management, improving application performance",
        "Deployed applications on AWS cloud platform, utilizing EC2, S3, and RDS services",
        "Containerized applications using Docker and orchestrated with Kubernetes for scalable deployments",
        "Implemented CI/CD pipelines using Git version control and Jenkins, automating testing and deployment processes",
        "Applied Test-Driven Development (TDD) practices, creating comprehensive test suites achieving 85% code coverage",
        "Collaborated with cross-functional teams to design and implement API contracts and system architecture"
    ]
    for achievement in achievements2:
        doc.add_paragraph(achievement, style="List Bullet")
    doc.add_paragraph("")  # Empty line
    
    # Job 3
    doc.add_heading("Junior Software Engineer", level=2)
    doc.add_paragraph("TechVentures | San Jose, CA | Jul 2016 - May 2017")
    doc.add_paragraph("Key Achievements:")
    achievements3 = [
        "Developed backend features for web applications using Java, Python, and Spring Framework",
        "Gained experience with PostgreSQL database design and query optimization",
        "Learned and applied Git version control best practices in collaborative development environment",
        "Participated in agile development processes, sprint planning, and code review sessions",
        "Fixed bugs, improved application performance, and wrote comprehensive unit and integration tests",
        "Assisted in API development and learned RESTful API design principles"
    ]
    for achievement in achievements3:
        doc.add_paragraph(achievement, style="List Bullet")
    doc.add_paragraph("")  # Empty line
    
    # Education
    doc.add_heading("Education", level=1)
    doc.add_heading("Bachelor of Science in Computer Science", level=2)
    doc.add_paragraph("UC San Diego | San Diego, CA | 2016")
    doc.add_paragraph("Relevant Coursework: Data Structures, Algorithms, Database Systems, Software Engineering")
    doc.add_paragraph("")  # Empty line
    
    # Certifications
    doc.add_heading("Certifications", level=1)
    doc.add_paragraph("AWS Certified Solutions Architect - Associate (2021)")
    doc.add_paragraph("Kubernetes Certified Application Developer (2022)")
    doc.add_paragraph("")  # Empty line
    
    # Open Source Contributions
    doc.add_heading("Open Source Contributions", level=1)
    doc.add_paragraph("Contributor to open-source projects on GitHub, including contributions to Python libraries and Go microservices frameworks")
    doc.add_paragraph("Maintained personal projects demonstrating expertise in RESTful API design, microservices, and system design patterns")
    
    # Save document
    doc.save(output_path)
    print(f"Created resume: {output_path}")


if __name__ == "__main__":
    # Create sample directory structure
    project_root = Path(__file__).parent.parent
    jobs_dir = project_root / "data" / "sample" / "jobs"
    resumes_dir = project_root / "data" / "sample" / "resumes"
    jobs_dir.mkdir(parents=True, exist_ok=True)
    resumes_dir.mkdir(parents=True, exist_ok=True)
    
    # Create job posting
    job_posting_path = jobs_dir / "raw_job_posting.docx"
    create_software_engineer_job_posting(job_posting_path)
    
    # Create resumes
    data_scientist_resume_path = resumes_dir / "raw_resume_data_scientist.docx"
    create_data_scientist_resume(data_scientist_resume_path)
    
    software_engineer_resume_path = resumes_dir / "raw_resume_software_engineer.docx"
    create_software_engineer_resume(software_engineer_resume_path)
    
    print("\n" + "="*60)
    print("Sample documents created successfully!")
    print("="*60)
    print(f"\nCreated files:")
    print(f"  Jobs: {jobs_dir}")
    print("    - raw_job_posting.docx (Software Engineer - Backend)")
    print(f"  Resumes: {resumes_dir}")
    print("    - raw_resume_data_scientist.docx (Sarah Chen - Data Scientist)")
    print("    - raw_resume_software_engineer.docx (Michael Rodriguez - Software Engineer)")
    print("\nYou can now use these files to test your HR recruiting pipeline!")
